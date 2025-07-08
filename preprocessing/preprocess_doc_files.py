import docx
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
import re
from typing import List, Dict, Optional, Tuple
import unicodedata
import time
from datetime import datetime, timezone
from pathlib import Path
import io


# Endnotes, headers, footers, non-body parts not included in the processing

def clean_text(text: str) -> str:
    """
    Clean and normalize text by removing extra whitespace, special characters,
    and normalizing unicode characters. Preserves table formatting.
    """
    # print("Starting text cleaning...")
    
    # Split text into table and non-table parts
    parts = text.split("\nTable Content:")
    cleaned_parts = []
    
    # Clean the first part (non-table content)
    if parts[0].strip():
        # print("Cleaning non-table content...")
        # Normalize unicode characters
        cleaned = unicodedata.normalize('NFKD', parts[0])
        # print("Unicode normalization completed")
        
        # Remove special characters but keep basic punctuation
        cleaned = re.sub(r'[^\w\s.,!?-]', ' ', cleaned)
        # print("Special characters removed")
        
        # Remove extra whitespace
        cleaned = re.sub(r'\s+', ' ', cleaned)
        # print("Extra whitespace removed")
        
        cleaned_parts.append(cleaned.strip())
    
    # Process remaining parts (tables)
    for part in parts[1:]:
        if part.strip():
            # print("Preserving table content...")
            # Add back the "Table Content:" marker
            cleaned_parts.append( part)
    
    return '\n'.join(cleaned_parts)

def process_table(table) -> str:
    """
    Process a table and return its content in a block (Markdown-like) format with clear start/end markers.
    """
    print(f"Processing table with {len(table.rows)} rows...")
    table_lines = []
    # Extract header row
    header_cells = [clean_text(cell.text) for cell in table.rows[0].cells]
    table_lines.append('| ' + ' | '.join(header_cells) + ' |')
    table_lines.append('|' + '|'.join(['---'] * len(header_cells)) + '|')
    # Extract data rows
    for row in table.rows[1:]:
        row_cells = [clean_text(cell.text) for cell in row.cells]
        table_lines.append('| ' + ' | '.join(row_cells) + ' |')
    # Wrap with markers, but only one Table Content:
    return f"Table Content:\n" + '\n'.join(table_lines) + "\nEnd of table content.\n"

def extract_text_from_docx(file) -> List[Tuple[str, str]]:
    """
    Extract text from a Word document, including tables but excluding images.
    Returns a list of tuples (content_type, content) to preserve order.
    """
    try:
        if isinstance(file["content"], bytes):
            file["content"] = Document(io.BytesIO(file["content"]))
        doc = file["content"]
        content_parts = []
        
        # print("\nProcessing document in order...")
        current_paragraphs = []
        
        # Process the document in order
        for element in doc.element.body:
            if element.tag.endswith('p'):  # Paragraph
                para = Paragraph(element, doc)
                # Skip paragraphs that only contain images
                if not any(run._element.xpath('.//w:drawing') for run in para.runs):
                    if para.text.strip():
                        current_paragraphs.append(para.text)
            
            elif element.tag.endswith('tbl'):  # Table
                # Add any accumulated paragraphs first
                if current_paragraphs:
                    content_parts.append(('text', '\n'.join(current_paragraphs)))
                    current_paragraphs = []
                
                # Process the table
                table = Table(element, doc)
                table_text = process_table(table)
                if table_text:
                    content_parts.append(('table', table_text))
        
        # Add any remaining paragraphs
        if current_paragraphs:
            content_parts.append(('text', '\n'.join(current_paragraphs)))
        
        # print("\nProcessing headers and footers...")
        # Process headers and footers
        for section_idx, section in enumerate(doc.sections):
            # print(f"Processing section {section_idx + 1}")
            for header in section.header.paragraphs:
                if header.text.strip():
                    content_parts.append(('header', f"Header: {header.text}"))
            for footer in section.footer.paragraphs:
                if footer.text.strip():
                    content_parts.append(('footer', f"Footer: {footer.text}"))
        
        # print("\nDocument processing completed")
        return content_parts
    except Exception as e:
        # print(f"\nError during document processing: {str(e)}")
        raise Exception(f"Error extracting text from document: {str(e)}")


def split_text_into_chunks(text: str, chunk_size: int = 1500, overlap: int = 200) -> List[str]:
    """
    Split text into overlapping chunks of specified size.
    - If a table starts in a chunk and cannot be stored as a whole, move it to the next chunk.
    - If a table is larger than the chunk size, split by rows, each chunk starts with the header and separator.
    - If there is no table in the chunk, use the original chunking logic (with your overlap logic).
    - If a table fits in the chunk, keep it with other data in the same chunk.
    """
    chunks = []
    text_length = len(text)
    start = 0

    min_overlap = max(50, int(overlap * 0.5)) # Ensure a minimum overlap, at least 50
    max_overlap = min(overlap * 1.5, chunk_size - 50) # Max overlap should not be too close to chunk_size

    # Find all tables
    table_matches = list(re.finditer(r"(Table Content:(?! Table Content:).*?End of table content\.)", text, re.DOTALL))
    table_spans = [m.span() for m in table_matches]

    while start < text_length:
        # Find the next table that starts after 'start'
        next_table = None
        for t_start, t_end in table_spans:
            if t_start >= start:
                next_table = (t_start, t_end)
                break

        end = start + chunk_size
        if end > text_length:
            end = text_length

        if next_table and next_table[0] < end:
            t_start, t_end = next_table
            # If the table fits entirely within the chunk, include it with other data
            if t_end <= end:
                chunk = text[start:end]
                chunks.append(chunk)
                print(f"Created chunk {len(chunks)} (length: {len(chunk)})")
                if end == text_length:
                    break
                # Overlap logic (your original)
                potential_overlap_start = max(start, end - overlap)
                overlap_candidate = -1
                for i in range(end - min_overlap, potential_overlap_start -1, -1):
                    if text[i] == '.' or text[i] == ' ' or text[i] == '\n':
                        overlap_candidate = i
                        break
                current_overlap = overlap
                if overlap_candidate != -1:
                    calculated_overlap = end - overlap_candidate
                    if min_overlap <= calculated_overlap <= max_overlap:
                        current_overlap = calculated_overlap
                    elif calculated_overlap < min_overlap:
                        current_overlap = min_overlap
                    else:
                        current_overlap = max_overlap
                else:
                    current_overlap = max(min_overlap, min(overlap, max_overlap))
                    print(f"No strategic overlap point found. Using default/clamped overlap: {current_overlap}")
                start = end - current_overlap
            else:
                # Table would be split, so move it to the next chunk
                if t_start > start:
                    # Add text before the table as a chunk (use original logic)
                    pre_table_text = text[start:t_start]
                    if pre_table_text.strip():
                        pre_chunks = _original_chunking(pre_table_text, chunk_size, overlap)
                        chunks.extend(pre_chunks)
                # Now handle the table
                table_text = text[t_start:t_end]
                table_lines = table_text.split('\n')
                # Find header and separator (assume Markdown format)
                header = table_lines[1] if len(table_lines) > 1 else ''
                separator = table_lines[2] if len(table_lines) > 2 else ''
                rows = table_lines[3:-1]  # skip header, separator, and last line ("End of table content.")
                current_chunk = f"Table Content:\n{table_lines[0]}\n{header}\n{separator}"
                for row in rows:
                    if len(current_chunk) + len(row) + 1 > chunk_size:
                        current_chunk += "\nEnd of table content."
                        chunks.append(current_chunk)
                        print(f"Created chunk {len(chunks)} (length: {len(current_chunk)})")
                        current_chunk = f"Table Content:\n{table_lines[0]}\n{header}\n{separator}\n{row}"
                    else:
                        current_chunk += f"\n{row}"
                if current_chunk.strip() and not current_chunk.strip().endswith("End of table content."):
                    current_chunk += "\nEnd of table content."
                    chunks.append(current_chunk)
                    print(f"Created chunk {len(chunks)} (length: {len(current_chunk)})")
                start = t_end
        else:
            # No table in this chunk, use original chunking logic
            end = start + chunk_size
            if end > text_length:
                end = text_length
            chunk = text[start:end]
            if chunk.strip():
                # Use your original chunking logic for this chunk
                pre_chunks = _original_chunking(chunk, chunk_size, overlap)
                chunks.extend(pre_chunks)
            if end == text_length:
                break
            # Overlap logic (your original)
            potential_overlap_start = max(start, end - overlap)
            overlap_candidate = -1
            for i in range(end - min_overlap, potential_overlap_start -1, -1):
                if text[i] == '.' or text[i] == ' ' or text[i] == '\n':
                    overlap_candidate = i
                    break
            current_overlap = overlap
            if overlap_candidate != -1:
                calculated_overlap = end - overlap_candidate
                if min_overlap <= calculated_overlap <= max_overlap:
                    current_overlap = calculated_overlap
                elif calculated_overlap < min_overlap:
                    current_overlap = min_overlap
                else:
                    current_overlap = max_overlap
            else:
                current_overlap = max(min_overlap, min(overlap, max_overlap))
                print(f"No strategic overlap point found. Using default/clamped overlap: {current_overlap}")
            start = end - current_overlap

    print(f"Text splitting completed. Total chunks: {len(chunks)}")
    return chunks

# Helper function for original chunking logic

def _original_chunking(text: str, chunk_size: int, overlap: int) -> List[str]:
    chunks = []
    text_length = len(text)
    start = 0
    min_overlap = max(50, int(overlap * 0.5)) # Ensure a minimum overlap, at least 50
    max_overlap = min(overlap * 1.5, chunk_size - 50) # Max overlap should not be too close to chunk_size
    while start < text_length:
        end = start + chunk_size
        if end < text_length:
            last_period = text.rfind('.', start, end)
            last_space = text.rfind(' ', start, end)
            last_newline = text.rfind('\n', start, end)
            split_point = max(last_period, last_space, last_newline)
            if split_point > start:
                end = split_point + 1
        else:
            end = text_length
        chunk = text[start:end]
        chunks.append(chunk)
        print(f"Created chunk {len(chunks)} (length: {len(chunk)})")
        if end == text_length:
            break
        potential_overlap_start = max(start, end - overlap)
        overlap_candidate = -1
        for i in range(end - min_overlap, potential_overlap_start -1, -1):
            if text[i] == '.' or text[i] == ' ' or text[i] == '\n':
                overlap_candidate = i
                break
        current_overlap = overlap
        if overlap_candidate != -1:
            calculated_overlap = end - overlap_candidate
            if min_overlap <= calculated_overlap <= max_overlap:
                current_overlap = calculated_overlap
            elif calculated_overlap < min_overlap:
                current_overlap = min_overlap
            else:
                current_overlap = max_overlap
        else:
            current_overlap = max(min_overlap, min(overlap, max_overlap))
            print(f"No strategic overlap point found. Using default/clamped overlap: {current_overlap}")
        start = end - current_overlap
    return chunks


def preprocess_doc(file = {}, chunk_size: int = 1000, overlap: int = 200, blob_metadata = None) -> Dict[str, List[str]]:
    """
    Main preprocessing function that handles the entire document preprocessing pipeline.
    """
    print(f"\nStarting document preprocessing: {file["name"]}")
    
    if isinstance(file["content"], bytes):
        file["content"] = Document(io.BytesIO(file["content"]))
    print("\nStep 1: Extracting text from document")
    content_parts = extract_text_from_docx(file)
    print(f"Extracted {len(content_parts)} content parts from document")
    # print("Content parts preview:", content_parts)
    
    # Process each part according to its type
    processed_parts = []
    for content_type, content in content_parts:
        if content_type == 'table':
            # Keep tables as is
            processed_parts.append(content)
        else:
            # Clean non-table content
            processed_parts.append(clean_text(content))
    
    # Join all parts together
    final_text = '\n'.join(processed_parts)
    print(f"Extracted {len(final_text)} characters of processed text")
    print('Processed text preview:', final_text[:500])
    
    print("\nStep 2: Splitting into chunks")
    chunks = split_text_into_chunks(final_text, chunk_size, overlap)
    timestamp = datetime.fromtimestamp(time.time(), tz=timezone.utc)
    # Prepare metadata
    metadata = {
        # TODO: userId, clientId, timestamp
        'user_id': 'example_user',
        'client_id': 'example_client',
        'timestamp': timestamp.strftime('%Y-%m-%dT%H:%M:%S.%f')[:-3] + 'Z',
        'num_chunks': len(chunks),
        'total_chars': len(final_text)
    }
    metadata.update(blob_metadata or {})
    metadata["has_been_preprocessed"] = "True"
    
    print("\nPreprocessing completed successfully!")
    return {
        'chunks': chunks,
        'metadata': metadata
    }


if __name__ == "__main__":
    # Example usage
    try:
        start_time = time.time()
        doc_path = r"test_data/sample.docx"
        print("\n=== Starting Document Processing ===")
        result = preprocess_doc(doc_path, blob_metadata=None)
        
        print("\n=== Processing Results ===")
        print(f"Processed document: {result['metadata']['filename']}")
        print(f"Number of chunks: {result['metadata']['num_chunks']}")
        print(f"Total characters: {result['metadata']['total_chars']}")
        print(f"Processing time: {time.time() - start_time:.2f} seconds")
        # Print first chunk as example
        if result['chunks']:
            for chunk in result['chunks']:
                print(f"\nChunk:\n{chunk}\n")
            
    except Exception as e:
        print(f"\nError processing document: {str(e)}")
