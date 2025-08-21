import docx
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
import re
from typing import List, Dict, Optional, Tuple
import unicodedata
import time
from datetime import datetime, timezone
from pathlib import Path
import io


def clean_text(text: str) -> str:
    """
    Clean and normalize text by removing extra whitespace, special characters,
    and normalizing unicode characters. Preserves table formatting.
    """
    parts = text.split("\nTable Content:")
    cleaned_parts = []
    if parts[0].strip():
        cleaned = unicodedata.normalize('NFKD', parts[0])
        cleaned = re.sub(r'[^\w\s.,!?-]', ' ', cleaned)
        cleaned = re.sub(r'\s+', ' ', cleaned)
        cleaned_parts.append(cleaned.strip())
    for part in parts[1:]:
        if part.strip():
            cleaned_parts.append( part)
    
    return '\n'.join(cleaned_parts)

def process_table(table) -> str:
    """
    Process a table and return its content in a block (Markdown-like) format with clear start/end markers.
    """
    print(f"Processing table with {len(table.rows)} rows...")
    table_lines = []
    # Extract header row
    header_cells = [clean_text(cell.text) for cell in table.rows[0].cells]
    table_lines.append('| ' + ' | '.join(header_cells) + ' |')
    table_lines.append('|' + '|'.join(['---'] * len(header_cells)) + '|')
    # Extract data rows
    for row in table.rows[1:]:
        row_cells = [clean_text(cell.text) for cell in row.cells]
        table_lines.append('| ' + ' | '.join(row_cells) + ' |')
    # Wrap with markers, but only one Table Content:
    return f"Table Content:\n" + '\n'.join(table_lines) + "\nEnd of table content.\n"

def extract_text_from_docx(file) -> List[Tuple[str, str]]:
    """
    Extract text from a Word document, including tables but excluding images.
    Returns a list of tuples (content_type, content) to preserve order.
    """
    try:
        if isinstance(file["content"], bytes):
            file["content"] = Document(io.BytesIO(file["content"]))
        doc = file["content"]
        content_parts = []
        
        # print("\nProcessing document in order...")
        current_paragraphs = []
        
        # Process the document in order
        for element in doc.element.body:
            if element.tag.endswith('p'):  # Paragraph
                para = Paragraph(element, doc)
                # Skip paragraphs that only contain images
                if not any(run._element.xpath('.//w:drawing') for run in para.runs):
                    if para.text.strip():
                        current_paragraphs.append(para.text)
            
            elif element.tag.endswith('tbl'):  # Table
                # Add any accumulated paragraphs first
                if current_paragraphs:
                    content_parts.append(('text', '\n'.join(current_paragraphs)))
                    current_paragraphs = []
                
                # Process the table
                table = Table(element, doc)
                table_text = process_table(table)
                if table_text:
                    content_parts.append(('table', table_text))
        
        # Add any remaining paragraphs
        if current_paragraphs:
            content_parts.append(('text', '\n'.join(current_paragraphs)))
        
        # print("\nProcessing headers and footers...")
        # Process headers and footers
        for section_idx, section in enumerate(doc.sections):
            # print(f"Processing section {section_idx + 1}")
            for header in section.header.paragraphs:
                if header.text.strip():
                    content_parts.append(('header', f"Header: {header.text}"))
            for footer in section.footer.paragraphs:
                if footer.text.strip():
                    content_parts.append(('footer', f"Footer: {footer.text}"))
        
        # print("\nDocument processing completed")
        return content_parts
    except Exception as e:
        # print(f"\nError during document processing: {str(e)}")
        raise Exception(f"Error extracting text from document: {str(e)}")

def split_text_into_chunks(text: str, chunk_size: int = 1500, overlap: int = 200) -> List[str]:
    """
    Split text into overlapping chunks of specified size.
    - If a table starts in a chunk and cannot be stored as a whole, move it to the next chunk.
    - If a table is larger than the chunk size, split by rows, each chunk starts with the header and separator.
    - If there is no table in the chunk, use the original chunking logic (with your overlap logic).
    - If a table fits in the chunk, keep it with other data in the same chunk.
    """
    chunks = []
    text_length = len(text)
    start = 0

    min_overlap = max(50, int(overlap * 0.5)) # Ensure a minimum overlap, at least 50
    max_overlap = min(overlap * 1.5, chunk_size - 50) # Max overlap should not be too close to chunk_size

    # Find all tables
    table_matches = list(re.finditer(r"(Table Content:(?! Table Content:).*?End of table content\.)", text, re.DOTALL))
    table_spans = [m.span() for m in table_matches]

    while start < text_length:
        # Find the next table that starts after 'start'
        next_table = None
        for t_start, t_end in table_spans:
            if t_start >= start:
                next_table = (t_start, t_end)
                break

        end = start + chunk_size
        if end > text_length:
            end = text_length

        if next_table and next_table[0] < end:
            t_start, t_end = next_table
            # If the table fits entirely within the chunk, include it with other data
            if t_end <= end:
                # Find proper sentence boundary before the table or at the end
                best_end = end
                if end < text_length:
                    # Look for sentence boundaries first (period, exclamation, question mark)
                    sentence_end = -1
                    for i in range(end, max(start + chunk_size // 2, start), -1):
                        if i < text_length and text[i-1] in '.!?' and (i == text_length or text[i].isspace() or text[i].isupper()):
                            sentence_end = i
                            break
                    
                    if sentence_end != -1:
                        best_end = sentence_end
                    else:
                        # If no sentence boundary found, fall back to other boundaries
                        for i in range(end, max(start + chunk_size // 2, start), -1):
                            if i < text_length and text[i-1] in '.\n':
                                best_end = i
                                break
                
                chunk = text[start:best_end]
                chunks.append(chunk)
                print(f"Created chunk {len(chunks)} (length: {len(chunk)})")
                if best_end >= text_length:
                    break
                
                # Find where to start the next chunk (after any whitespace)
                next_start = best_end
                while next_start < text_length and text[next_start].isspace():
                    next_start += 1
                
                start = next_start
            else:
                # Table would be split, so move it to the next chunk
                if t_start > start:
                    # Add text before the table as a chunk (use original logic)
                    pre_table_text = text[start:t_start]
                    if pre_table_text.strip():
                        pre_chunks = _original_chunking(pre_table_text, chunk_size, overlap)
                        chunks.extend(pre_chunks)
                # Now handle the table
                table_text = text[t_start:t_end]
                table_lines = table_text.split('\n')
                # Find header and separator (assume Markdown format)
                header = table_lines[1] if len(table_lines) > 1 else ''
                separator = table_lines[2] if len(table_lines) > 2 else ''
                rows = table_lines[3:-1]  # skip header, separator, and last line ("End of table content.")
                current_chunk = f"Table Content:\n{table_lines[0]}\n{header}\n{separator}"
                for row in rows:
                    if len(current_chunk) + len(row) + 1 > chunk_size:
                        current_chunk += "\nEnd of table content."
                        chunks.append(current_chunk)
                        print(f"Created chunk {len(chunks)} (length: {len(current_chunk)})")
                        current_chunk = f"Table Content:\n{table_lines[0]}\n{header}\n{separator}\n{row}"
                    else:
                        current_chunk += f"\n{row}"
                if current_chunk.strip() and not current_chunk.strip().endswith("End of table content."):
                    current_chunk += "\nEnd of table content."
                    chunks.append(current_chunk)
                    print(f"Created chunk {len(chunks)} (length: {len(current_chunk)})")
                start = t_end
        else:
            # No table in this chunk, find proper sentence boundaries
            end = start + chunk_size
            if end >= text_length:
                end = text_length
                chunk = text[start:end]
                chunks.append(chunk)
                break
            else:
                # Find the best place to end the chunk at a sentence boundary
                best_end = end
                # Look backwards from the target end to find a sentence ending
                sentence_end = -1
                for i in range(end, max(start + chunk_size // 2, start), -1):
                    if i < text_length and text[i-1] in '.!?' and (i == text_length or text[i].isspace() or text[i].isupper()):
                        sentence_end = i
                        break
                
                if sentence_end != -1:
                    best_end = sentence_end
                else:
                    # If no sentence boundary found, look for other boundaries
                    for i in range(end, max(start + chunk_size // 2, start), -1):
                        if i < text_length and text[i-1] in '.\n':
                            best_end = i
                            break
                
                chunk = text[start:best_end]
                chunks.append(chunk)
                print(f"Created chunk {len(chunks)} (length: {len(chunk)})")
                
                # Find where to start the next chunk (after any whitespace)
                next_start = best_end
                while next_start < text_length and text[next_start].isspace():
                    next_start += 1
                
                start = next_start

    print(f"Text splitting completed. Total chunks: {len(chunks)}")
    return chunks

def _original_chunking(text: str, chunk_size: int, overlap: int) -> List[str]:
    chunks = []
    text_length = len(text)
    start = 0
    min_overlap = max(50, int(overlap * 0.5)) # Ensure a minimum overlap, at least 50
    max_overlap = min(overlap * 1.5, chunk_size - 50) # Max overlap should not be too close to chunk_size
    while start < text_length:
        end = start + chunk_size
        if end < text_length:
            # Look for sentence boundaries first (period, exclamation, question mark)
            sentence_end = -1
            for i in range(end, start, -1):
                if text[i-1] in '.!?' and i < len(text) and (text[i].isspace() or text[i].isupper()):
                    sentence_end = i
                    break
            
            # If no sentence boundary found, fall back to other boundaries
            if sentence_end == -1:
                last_period = text.rfind('.', start, end)
                last_exclamation = text.rfind('!', start, end)
                last_question = text.rfind('?', start, end)
                last_newline = text.rfind('\n', start, end)
                last_space = text.rfind(' ', start, end)
                split_point = max(last_period, last_exclamation, last_question, last_newline, last_space)
            else:
                split_point = sentence_end
            if split_point > start:
                end = split_point + 1
        else:
            end = text_length
        chunk = text[start:end]
        chunks.append(chunk)
        print(f"Created chunk {len(chunks)} (length: {len(chunk)})")
        if end == text_length:
            break
        potential_overlap_start = max(start, end - overlap)
        overlap_candidate = -1
        for i in range(end - min_overlap, potential_overlap_start -1, -1):
            if text[i] in '.!?' and i + 1 < len(text) and (text[i + 1].isspace() or text[i + 1].isupper()):
                overlap_candidate = i + 1  # Start after the sentence end
                break
            elif text[i] == '.' or text[i] == ' ' or text[i] == '\n':
                overlap_candidate = i + 1
                break
        current_overlap = overlap
        if overlap_candidate != -1:
            calculated_overlap = end - overlap_candidate
            if min_overlap <= calculated_overlap <= max_overlap:
                current_overlap = calculated_overlap
            elif calculated_overlap < min_overlap:
                current_overlap = min_overlap
            else:
                current_overlap = max_overlap
        else:
            current_overlap = max(min_overlap, min(overlap, max_overlap))
            print(f"No strategic overlap point found. Using default/clamped overlap: {current_overlap}")
        start = end - current_overlap
    return chunks

def preprocess_doc(file = {}, chunk_size: int = 1000, overlap: int = 200, blob_metadata = None) -> Dict[str, List[str]]:
    """
    Main preprocessing function that handles the entire document preprocessing pipeline.
    """
    print(f"\nStarting document preprocessing: {file["name"]}")
    
    if isinstance(file["content"], bytes):
        file["content"] = Document(io.BytesIO(file["content"]))
    print("\nStep 1: Extracting text from document")
    content_parts = extract_text_from_docx(file)
    print(f"Extracted {len(content_parts)} content parts from document")
    # print("Content parts preview:", content_parts)
    
    # Process each part according to its type
    processed_parts = []
    for content_type, content in content_parts:
        if content_type == 'table':
            # Keep tables as is
            processed_parts.append(content)
        else:
            # Clean non-table content
            processed_parts.append(clean_text(content))
    
    # Join all parts together
    final_text = '\n'.join(processed_parts)
    print(f"Extracted {len(final_text)} characters of processed text")
    print('Processed text preview:', final_text[:500])
    
    print("\nStep 2: Splitting into chunks")
    chunks = split_text_into_chunks(final_text, chunk_size, overlap)
    
    # Normalize chunks
    normalized_chunks = []
    for chunk in chunks:
        normalized_chunk = normalize_chunk(chunk)
        normalized_chunks.append(normalized_chunk)
        print(f"Normalized chunk: {normalized_chunk[:100]}...")
    
    print(f"Chunk length for document {file['name']} = {len(normalized_chunks)}")
    
    timestamp = datetime.fromtimestamp(time.time(), tz=timezone.utc)
    # Prepare metadata
    metadata = {
        # TODO: userId, clientId, timestamp
        'user_id': 'example_user',
        'client_id': 'example_client',
        'timestamp': timestamp.strftime('%Y-%m-%dT%H:%M:%S.%f')[:-3] + 'Z',
        'num_chunks': len(normalized_chunks),
        'total_chars': len(final_text)
    }
    metadata.update(blob_metadata or {})
    metadata["has_been_preprocessed"] = "True"
    
    print("\nPreprocessing completed successfully!")
    return {
        'result': {
            'chunks': normalized_chunks,
            'metadata': metadata
        }
    }

def normalize_chunk(chunk: str) -> str:
    """
    Normalize a text chunk by cleaning up formatting, spacing, and punctuation.
    """
    normalized_response = re.sub(r"\[\[\d+\]\],?\s*", "", chunk)
    normalized_response = re.sub(r"\n\s*#{1,6}\s*([^\n]+)", r". \1:", normalized_response)
    normalized_response = re.sub(r"\n\s*-\s*\*\*([^*]+)\*\*:\s*", r". \1: ", normalized_response)
    normalized_response = re.sub(r"\n\s*-\s*", ". ", normalized_response)
    normalized_response = re.sub(r"\*\*([^*]+)\*\*", r"\1", normalized_response)
    normalized_response = re.sub(r"\*([^*]+)\*", r"\1", normalized_response)
    normalized_response = re.sub(r"\n{3,}", "\n\n", normalized_response)
    normalized_response = re.sub(r"\n\n", ". ", normalized_response)
    normalized_response = re.sub(r"\n", " ", normalized_response)
    
    # Handle hyphenated words broken across lines
    normalized_response = re.sub(r"(\w)-\s+(\w)", r"\1\2", normalized_response)
    
    # Remove extra hyphens and normalize dashes
    normalized_response = re.sub(r"--+", "—", normalized_response)
    
    # Fix common document extraction artifacts
    normalized_response = re.sub(r"\s*\|\s*", " ", normalized_response)  # Remove table separators
    normalized_response = re.sub(r"(\w)\s*_\s*(\w)", r"\1_\2", normalized_response)  # Fix broken underscores
    
    # Handle page numbers and headers/footers patterns
    normalized_response = re.sub(r"\b(Page\s+\d+|©\s*\d+|All rights reserved)\b", "", normalized_response, flags=re.IGNORECASE)
    
    # Fix spacing around parentheses and brackets
    normalized_response = re.sub(r"\s*\(\s*", " (", normalized_response)
    normalized_response = re.sub(r"\s*\)\s*", ") ", normalized_response)
    normalized_response = re.sub(r"\s*\[\s*", " [", normalized_response)
    normalized_response = re.sub(r"\s*\]\s*", "] ", normalized_response)
    
    # Clean up punctuation
    normalized_response = re.sub(r"\.+", ".", normalized_response)
    normalized_response = re.sub(r"\s*\.\s*\.", ".", normalized_response)
    normalized_response = re.sub(r":\s*\.", ":", normalized_response)
    normalized_response = re.sub(r"\.\s*:", ":", normalized_response)
    
    # Clean up multiple punctuation marks
    normalized_response = re.sub(r"[.]{2,}", ".", normalized_response)
    normalized_response = re.sub(r"[,]{2,}", ",", normalized_response)
    
    normalized_response = re.sub(r"\s+", " ", normalized_response)
    normalized_response = re.sub(r"\s*([.,:;!?])", r"\1", normalized_response)
    normalized_response = re.sub(r"([.,:;!?])\s*", r"\1 ", normalized_response)
    
    # Fix spacing before opening quotes and after closing quotes
    normalized_response = re.sub(r'\s*"([^"]*)"\s*', r' "\1" ', normalized_response)
    normalized_response = re.sub(r"\s*'([^']*)'\s*", r" '\1' ", normalized_response)
    
    normalized_response = re.sub(r"\.\s*([a-z])", lambda m: ". " + m.group(1).upper(), normalized_response)
    normalized_response = re.sub(r"(\d+)\.\s+(\d+)", r"\1.\2", normalized_response)
    
    # Final cleanup - remove any remaining double spaces
    normalized_response = re.sub(r"\s{2,}", " ", normalized_response)
    
    normalized_response = normalized_response.strip()
    if normalized_response and not normalized_response[0].isupper():
        normalized_response = normalized_response[0].upper() + normalized_response[1:]
    if normalized_response and normalized_response[-1] not in '.!?':
        normalized_response += "."
    return normalized_response

if __name__ == "__main__":
    try:
        start_time = time.time()
        doc_path = r"test_data/sample.docx"
        print("\n=== Starting Document Processing ===")
        
        with open(doc_path, 'rb') as f:
            file_content = f.read()
        
        file_dict = {"name": doc_path, "content": file_content}
        result = preprocess_doc(file_dict, blob_metadata=None)
        
        print("\n=== Processing Results ===")
        print(f"Processed document: {result['result']['metadata']['user_id']}")
        print(f"Number of chunks: {result['result']['metadata']['num_chunks']}")
        print(f"Total characters: {result['result']['metadata']['total_chars']}")
        print(f"Processing time: {time.time() - start_time:.2f} seconds")
        
        if result['result']['chunks']:
            print(f"\nFirst chunk example:\n{result['result']['chunks'][0]}\n")
            
    except Exception as e:
        print(f"\nError processing document: {str(e)}")
